"""
APA 7 Formatting: python-docx for .docx (double spacing, 1" margins, Times New Roman 12pt,
running head, page numbers, hanging indents for references). Optional PDF via reportlab.
"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from articlewriter.models import ArticleSections
from articlewriter.utils import xml_escape


def _set_cell_margin(cell: Any, margin_inch: float) -> None:
    """Set table cell margin (used for spacing)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in [("top", margin_inch), ("left", margin_inch), ("bottom", margin_inch), ("right", margin_inch)]:
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(int(val * 1440)))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


class APAFormatter:
    """
    Produces APA 7 formatted .docx and optional .pdf. Double spacing, 1" margins,
    Times New Roman 12pt, running head, page numbers, hanging indent references.
    """

    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
        margin_inches: float = 1.0,
        line_spacing: float = 2.0,
        running_head: bool = True,
        add_disclaimer: bool = False,
    ):
        self.font_name = font_name
        self.font_size = font_size
        self.margin_inches = margin_inches
        self.line_spacing = line_spacing
        self.running_head = running_head
        self.add_disclaimer = add_disclaimer

    def _style_paragraph(self, p: Any) -> None:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = self.line_spacing
        for r in p.runs:
            r.font.name = self.font_name
            r.font.size = Pt(self.font_size)

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = self.line_spacing

    def _add_body(self, doc: Document, text: str) -> None:
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = self.line_spacing
        run = p.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)

    def _add_reference(self, doc: Document, apa_string: str, hanging: bool = True) -> None:
        p = doc.add_paragraph()
        if hanging:
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = self.line_spacing
        run = p.add_run(apa_string)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)

    def to_docx(self, sections: ArticleSections, path: str | Path) -> Path:
        """Write full manuscript to .docx with APA 7 formatting."""
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(self.margin_inches)
        sec.bottom_margin = Inches(self.margin_inches)
        sec.left_margin = Inches(self.margin_inches)
        sec.right_margin = Inches(self.margin_inches)

        # Title (centered)
        self._add_heading(doc, sections.title, level=0)
        doc.add_paragraph()

        # Running head (simplified: short title left, page number right)
        if self.running_head:
            head = doc.add_paragraph()
            head.paragraph_format.space_after = Pt(0)
            run_left = head.add_run(sections.title[:50] + ("..." if len(sections.title) > 50 else ""))
            run_left.font.name = self.font_name
            run_left.font.size = Pt(12)
            run_right = head.add_run("\t1")
            run_right.font.name = self.font_name
            run_right.font.size = Pt(12)
            head.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()

        # Abstract
        self._add_heading(doc, "Abstract", level=1)
        self._add_body(doc, sections.abstract)
        doc.add_paragraph()

        # Keywords
        if sections.keywords:
            kw = doc.add_paragraph()
            run = kw.add_run("Keywords: " + ", ".join(sections.keywords))
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run.italic = True
            kw.paragraph_format.line_spacing = self.line_spacing
            doc.add_paragraph()

        # Main sections
        for heading, body in [
            ("Introduction", sections.introduction),
            ("Literature Review", sections.literature_review),
            ("Methodology", sections.methodology),
            ("Results", sections.results),
            ("Discussion", sections.discussion),
            ("Theoretical Implications", sections.theoretical_implications),
            ("Practical Implications", sections.practical_implications),
            ("Limitations", sections.limitations),
            ("Future Research", sections.future_research),
            ("Conclusion", sections.conclusion),
        ]:
            if body:
                self._add_heading(doc, heading, level=1)
                self._add_body(doc, body)
                doc.add_paragraph()

        # References (hanging indent)
        self._add_heading(doc, "References", level=1)
        for ref in sections.references:
            apa = ref.get("apa_string") if isinstance(ref, dict) else str(ref)
            if apa:
                self._add_reference(doc, apa, hanging=True)

        # Optional academic integrity disclaimer (footer note)
        if getattr(self, "add_disclaimer", False):
            p = doc.add_paragraph()
            run = p.add_run(
                "Disclaimer: This draft was generated with AI-assisted tools for research support. "
                "Authors are responsible for verification, originality, and compliance with journal and institutional policies."
            )
            run.font.name = self.font_name
            run.font.size = Pt(10)
            run.italic = True

        doc.save(str(path))
        return path

    def to_pdf(
        self,
        sections: ArticleSections,
        pdf_path: str | Path,
        docx_path: str | Path | None = None,
    ) -> Path:
        """
        Produce PDF: convert from .docx if docx_path given and docx2pdf available;
        otherwise build PDF from sections using reportlab.
        """
        pdf_path = Path(pdf_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        if docx_path:
            try:
                from docx2pdf import convert as docx2pdf_convert
                docx2pdf_convert(str(docx_path), str(pdf_path))
                return pdf_path
            except ImportError:
                pass
        # Fallback: build PDF with reportlab from sections
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            leftMargin=inch * self.margin_inches,
            rightMargin=inch * self.margin_inches,
            topMargin=inch * self.margin_inches,
            bottomMargin=inch * self.margin_inches,
        )
        styles = getSampleStyleSheet()
        normal = ParagraphStyle(
            "CustomNormal",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=self.font_size,
            leading=self.font_size * self.line_spacing,
            spaceAfter=12,
            leftIndent=36,
            firstLineIndent=36,
        )
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontName="Times-Bold",
            fontSize=self.font_size,
            alignment=1,
            spaceAfter=24,
        )
        story = []
        story.append(Paragraph(xml_escape(sections.title), title_style))
        story.append(Spacer(1, 24))
        for heading, body in [
            ("Abstract", sections.abstract),
            ("Introduction", sections.introduction),
            ("Literature Review", sections.literature_review),
            ("Methodology", sections.methodology),
            ("Results", sections.results),
            ("Discussion", sections.discussion),
            ("Theoretical Implications", sections.theoretical_implications),
            ("Practical Implications", sections.practical_implications),
            ("Limitations", sections.limitations),
            ("Future Research", sections.future_research),
            ("Conclusion", sections.conclusion),
        ]:
            if body:
                story.append(Paragraph(f"<b>{heading}</b>", normal))
                story.append(Paragraph(xml_escape(body).replace("\n", "<br/>"), normal))
        story.append(Paragraph("<b>References</b>", normal))
        for ref in sections.references:
            apa = ref.get("apa_string", str(ref)) if isinstance(ref, dict) else str(ref)
            story.append(Paragraph(xml_escape(apa).replace("\n", " "), normal))
        doc.build(story)
        return pdf_path
